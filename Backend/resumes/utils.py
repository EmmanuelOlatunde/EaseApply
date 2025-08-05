import logging
import re
from typing import Dict, List, Optional, Any
from django.core.files.uploadedfile import UploadedFile


try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass

class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors"""
    pass

def extract_text_from_pdf(file) -> str:
    """
    Extract text from PDF file using PyPDF2
    
    Args:
        file: Django UploadedFile or file-like object
        
    Returns:
        str: Extracted text content
        
    Raises:
        TextExtractionError: If extraction fails
    """
    if PyPDF2 is None:
        raise TextExtractionError("PyPDF2 is not installed. Install with: pip install PyPDF2")
    
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        pdf_reader = PyPDF2.PdfReader(file)
        text_content = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                continue
        
        if not text_content:
            raise TextExtractionError("No text could be extracted from the PDF")
            
        return '\n\n'.join(text_content)
        
    except Exception as e:
        logger.error(f"PDF text extraction failed: {str(e)}")
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")
    finally:
        # Reset file pointer
        file.seek(0)

def extract_text_from_docx(file) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        file: Django UploadedFile or file-like object
        
    Returns:
        str: Extracted text content
        
    Raises:
        TextExtractionError: If extraction fails
    """
    if Document is None:
        raise TextExtractionError("python-docx is not installed. Install with: pip install python-docx")
    
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        doc = Document(file)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        if not text_content:
            raise TextExtractionError("No text could be extracted from the DOCX file")
            
        return '\n\n'.join(text_content)
        
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")
    finally:
        # Reset file pointer
        file.seek(0)

def extract_text_from_resume(file, file_type: str) -> str:
    """
    Extract text from resume file based on file type
    
    Args:
        file: Django UploadedFile or file-like object
        file_type: Either 'pdf' or 'docx'
        
    Returns:
        str: Extracted text content
        
    Raises:
        TextExtractionError: If extraction fails or unsupported file type
    """
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file)
    else:
        raise TextExtractionError(f"Unsupported file type: {file_type}")

def validate_resume_file(file: UploadedFile) -> tuple[str, bool]:
    """
    Validate uploaded resume file
    
    Args:
        file: Django UploadedFile
        
    Returns:
        tuple: (file_type, is_valid)
    """
    if not file:
        return '', False
    
    # Get file extension
    filename = file.name.lower()
    file_type = ''
    if filename.endswith('.pdf'):
        file_type = 'pdf'
    elif filename.endswith(('.docx', '.doc')):
        file_type = 'docx'

    # Check file size (limit to 10MB)
    max_size = 10 * 1024 * 1024  # 10MB
    if file.size >= max_size:
        return file_type, False

    return file_type, bool(file_type)  # Valid if file_type is non-empty

def extract_full_name(text: str) -> Optional[str]:
    """
    Extract full name from resume text
    """
    job_title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'consultant', 'specialist']
    lines = text.strip().split('\n')
    for line in lines[:10]:
        line_clean = line.strip().lower()
        if line_clean.startswith('summary'):
            continue
        words = line.strip().split()
        # Check if line is likely a job title
        if any(keyword in line_clean for keyword in job_title_keywords):
            continue
        # Allow names with titles like Dr., Mr., Ms.
        if 2 <= len(words) <= 4 and all(w.replace('.', '').replace(',', '').isalpha() for w in words):
            return line.strip()
    return None

def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
    """
    Extract contact information from resume text
    """
    contact_info = {
        'email': None,
        'phone': None,
        'address': None,
        'linkedin': None,
        'github': None
    }
    
    # Email extraction
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact_info['email'] = email_match.group()
    
    # Phone extraction
    phone_pattern = re.compile(
        r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    )
    phone_match = phone_pattern.search(text)
    if phone_match:
        contact_info['phone'] = phone_match.group()

        
    # LinkedIn extraction
    linkedin_pattern = r'(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9\-]+'
    linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_match:
        contact_info['linkedin'] = linkedin_match.group()
    
    # GitHub extraction
    github_pattern = r'(https?://)?(www\.)?github\.com/[A-Za-z0-9\-]+'
    github_match = re.search(github_pattern, text, re.IGNORECASE)
    if github_match:
        contact_info['github'] = github_match.group()
    
    # Address extraction (basic)
    lines = text.split('\n')
    for line in lines[:10]:  # Check first 10 lines
        if re.search(r'\b\d{5}(-\d{4})?\b', line):  # ZIP code pattern
            contact_info['address'] = line.strip()
            break
    
    return contact_info

def extract_summary(text: str) -> Optional[str]:
    """
    Extract summary from resume text - FIXED VERSION
    """
    summary_keywords = [
        'summary', 'professional summary', 'executive summary',
        'profile', 'professional profile', 'career summary',
        'objective', 'career objective', 'professional objective'
    ]
    
    stop_keywords = [
        'experience', 'professional experience', 'education', 'skills', 
        'employment', 'work history', 'projects', 'certifications', 'technical skills'
    ]
    
    lines = text.split('\n')
    summary_started = False
    summary_lines = []
    
    for i, line in enumerate(lines):
        line_clean = line.strip()
        line_for_match = line_clean.lower().replace(':', '')
        
        if not line_clean:
            if summary_started and summary_lines:
                # Empty line might indicate end of summary
                next_line = lines[i+1].strip().lower() if i+1 < len(lines) else ""
                if any(keyword in next_line for keyword in stop_keywords):
                    break
            continue
        
        if summary_started:
            # Check if this line is a section header
            if line_for_match in stop_keywords:
                break
            summary_lines.append(line_clean)
        elif any(keyword in line_for_match for keyword in summary_keywords):
            summary_started = True
            # Check if summary text is on the same line as header
            colon_split = line_clean.split(':', 1)
            if len(colon_split) > 1 and colon_split[1].strip():
                summary_lines.append(colon_split[1].strip())
    
    return ' '.join(summary_lines) if summary_lines else None

def extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text
    """
    skills = []
    skills_keywords = ['skills', 'technical skills', 'core competencies', 'technologies']
    
    lines = text.split('\n')
    skills_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        if ':' in line:
            parts = line.split(':', 1)[1]
            skills += [s.strip() for s in re.split(r'[,\|]', parts)]

        # Check if we've entered skills section
        if any(keyword in line_lower for keyword in skills_keywords):
            skills_section = True
            continue
        
        if skills_section:
            # Stop if we hit another major section
            if any(section in line_lower for section in [
                'experience', 'education', 'employment', 'work history',
                'projects', 'certifications', 'summary', 'objective'
            ]):
                break
            
            if line.strip():
                # Split by common delimiters
                line_skills = re.split(r'[,•·|]', line)
                for skill in line_skills:
                    skill = skill.strip()
                    if skill and len(skill) > 1:
                        skills.append(skill)
    
    return skills


def extract_work_experience(text: str) -> List[Dict[str, Any]]:
    """
    Extract work experience entries from resume.
    """
    experience = []
    lines = text.split('\n')
    
    exp_section = False
    current_job = {}
    
    # Flexible matching to identify experience section
    for idx, line in enumerate(lines):
        line_lower = line.lower().strip()
        original_line = line.strip()
        
        # Detect the start of experience section (more flexible matching)
        if not exp_section and ('experience' in line_lower and 
                               (line_lower.startswith('professional') or 
                                line_lower.startswith('work') or 
                                line_lower == 'experience' or
                                'professional experience' in line_lower)):
            exp_section = True
            continue

        if exp_section:
            # Detect the end of experience section
            if any(section in line_lower for section in [
                'education', 'skills', 'projects', 'certifications', 'summary'
            ]) and not any(word in line_lower for word in ['experience', 'work']):
                if current_job:
                    experience.append(current_job)
                    current_job = {}
                break

            if original_line:
                # Check if this looks like a job header
                # Look for patterns like: "Title | Company | Date" or "Title at Company (Date)"
                has_year = bool(re.search(r'\b(19|20)\d{2}\b', original_line))
                has_pipe = '|' in original_line
                
                # If it has a year and doesn't start with bullet points, it's likely a job header
                if has_year and not original_line.startswith(('•', '-', '*')):
                    # Save previous job if exists
                    if current_job:
                        experience.append(current_job)
                    
                    # Initialize new job
                    current_job = {
                        'title': '',
                        'company': '',
                        'duration': '',
                        'description': []
                    }
                    
                    if has_pipe:
                        # Handle pipe-separated format: "Title | Company | Date"
                        parts = [part.strip() for part in original_line.split('|')]
                        if len(parts) >= 3:
                            current_job['title'] = parts[0]
                            current_job['company'] = parts[1]
                            current_job['duration'] = parts[2]
                        elif len(parts) == 2:
                            current_job['title'] = parts[0]
                            current_job['company'] = parts[1]
                    else:
                        # Handle other formats
                        # Try to extract year range for duration
                        year_match = re.search(r'\b(19|20)\d{2}(-\d{4}|\s*-\s*(19|20)\d{2}|\s*-\s*present)?\b', original_line, re.IGNORECASE)
                        if year_match:
                            current_job['duration'] = year_match.group().strip()
                            # Remove the duration from the line to parse title and company
                            remaining = original_line.replace(year_match.group(), '').strip()
                            
                            # Try to split by common separators
                            if ' at ' in remaining:
                                parts = remaining.split(' at ', 1)
                                current_job['title'] = parts[0].strip()
                                current_job['company'] = parts[1].strip()
                            elif ' | ' in remaining:
                                parts = remaining.split(' | ', 1)
                                current_job['title'] = parts[0].strip()
                                current_job['company'] = parts[1].strip()
                            elif ' - ' in remaining:
                                parts = remaining.split(' - ', 1)
                                current_job['title'] = parts[0].strip()
                                current_job['company'] = parts[1].strip()
                            else:
                                # If we can't split, assume the whole thing is the title
                                current_job['title'] = remaining.strip()
                        else:
                            current_job['title'] = original_line

                # Handle bullet-point descriptions
                elif current_job and (original_line.startswith('•') or 
                                    original_line.startswith('-') or 
                                    original_line.startswith('*')):
                    current_job['description'].append(original_line)

    # Don't forget the last job
    if current_job:
        experience.append(current_job)

    return experience


def extract_education(text: str) -> List[Dict[str, str]]:
    """
    Extract education information from resume
    """
    education = []
    edu_keywords = ['education', 'academic background', 'qualifications']
    
    lines = text.split('\n')
    edu_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if we've entered education section
        if any(keyword in line_lower for keyword in edu_keywords):
            edu_section = True
            continue
        
        if edu_section:
            # Stop if we hit another major section
            if any(section in line_lower for section in [
                'experience', 'skills', 'projects', 'certifications'
            ]):
                break
            
            if line.strip() and re.search(r'\d{4}', line):  # Contains year
                education.append({
                    'degree': line.strip(),
                    'institution': '',
                    'year': re.search(r'\d{4}', line).group() if re.search(r'\d{4}', line) else ''
                })
    
    return education

def extract_certifications(text: str) -> List[str]:
    """
    Extract certifications from resume
    """
    certifications = []
    cert_keywords = ['certifications', 'certificates', 'professional certifications']
    
    lines = text.split('\n')
    cert_section = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if we've entered certifications section
        if any(keyword in line_lower for keyword in cert_keywords):
            cert_section = True
            continue
        
        if cert_section:
            # Stop if we hit another major section
            if any(section in line_lower for section in [
                'experience', 'education', 'skills', 'projects'
            ]):
                break
            
            if line.strip():
                certifications.append(line.strip())
    
    return certifications

def extract_projects(text: str) -> List[Dict[str, str]]:
    """
    Extract projects from resume
    """
    projects = []
    project_keywords = ['projects', 'personal projects', 'key projects']
    
    lines = text.split('\n')
    project_section = False
    current_project = {}
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if we've entered projects section
        if any(keyword in line_lower for keyword in project_keywords):
            project_section = True
            continue
        
        if project_section:
            # Stop if we hit another major section
            if any(section in line_lower for section in [
                'experience', 'education', 'skills', 'certifications'
            ]):
                if current_project:
                    projects.append(current_project)
                break
            
            if line.strip():
                if line.startswith('•') or line.startswith('-'):
                    if current_project:
                        if 'description' not in current_project:
                            current_project['description'] = []
                        current_project['description'].append(line.strip())
                else:
                    if current_project:
                        projects.append(current_project)
                    current_project = {
                        'name': line.strip(),
                        'description': []
                    }
    
    if current_project:
        projects.append(current_project)
    
    return projects

def parse_resume_content(extracted_text: str) -> Dict[str, Any]:
    """
    Parse extracted resume text into structured data
    
    Args:
        extracted_text: Raw text extracted from resume file
        
    Returns:
        Dict containing parsed resume information
        
    Raises:
        ResumeParsingError: If parsing fails
    """
    try:
        parsed_data = {
            'fullName': extract_full_name(extracted_text),
            'summary': extract_summary(extracted_text),
            'contactInfo': extract_contact_info(extracted_text),
            'skills': extract_skills(extracted_text),
            'workExperience': extract_work_experience(extracted_text),
            'education': extract_education(extracted_text),
            'certifications': extract_certifications(extracted_text),
            'projects': extract_projects(extracted_text)
        }
        
        logger.info("Successfully parsed resume content")
        return parsed_data
        
    except Exception as e:

        logger.error(f"Resume parsing failed: {str(e)}")
        raise ResumeParsingError(f"Failed to parse resume content: {str(e)}")
    

